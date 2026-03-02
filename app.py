import streamlit as st
import assemblyai as aai
import google.generativeai as genai
from docx import Document
import os
import tempfile
import time
import datetime

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Gestor de Actas Nilo (Final)", layout="wide")

st.title("🏛️ Prototipo de Actas - Concejo de Nilo")
st.markdown("Sistema Híbrido: **Texto Masivo** + **Audio** + **Recuperación**.")

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("🔑 Configuración")
    api_assembly = st.text_input("AssemblyAI API Key", type="password")
    api_google = st.text_input("Google Gemini API Key", type="password")
    
    # Selector de modelo (Automático)
    model_options = []
    index_default = 0
    if api_google:
        try:
            genai.configure(api_key=api_google)
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    model_options.append(m.name)
            # Intentar seleccionar flash automáticamente
            for i, m in enumerate(model_options):
                if "flash" in m:
                    index_default = i
                    break
        except:
            pass
    
    if model_options:
        selected_model = st.selectbox("Modelo:", model_options, index=index_default)
    else:
        selected_model = "models/gemini-1.5-flash"

# --- CONFIGURACIÓN DE SEGURIDAD (ANTIBLOQUEO) ---
safety_settings_nilo = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
]

# --- FUNCIONES LÓGICAS ---

def redactar_bloque_gemini(texto_crudo, orador_nombre, cargo, google_key, modelo_elegido):
    genai.configure(api_key=google_key)
    
    generation_config = {
        "temperature": 0.1, 
        "top_p": 0.95,
        "max_output_tokens": 8192,
    }

    try:
        model = genai.GenerativeModel(
            model_name=modelo_elegido, 
            generation_config=generation_config
        )
        
        info_orador = f"- Nombre: {orador_nombre}\n- Cargo: {cargo}" if orador_nombre else "- Orador: No especificado"

        prompt_usuario = f"""
        TU TAREA:
        Reescribir el siguiente texto oral pasando de PRIMERA PERSONA ("Yo/Nosotros") a TERCERA PERSONA ("El Concejal/La Administración/Ellos"), pero manteniendo la LITERALIDAD del discurso.
        
        DATOS DEL ORADOR:
        {info_orador}
        
        REGLAS DE ORO (ESTRICTAS):
        1. NO RESUMAS NADA: Si el texto repite palabras ("aprueban que aprueban"), DÉJALAS.
        2. NO USES LENGUAJE SOFISTICADO: No cambies "dijo" por "manifestó". Usa el vocabulario original.
        3. CONSERVACIÓN: Mantén expresiones como "¿Cierto?", "bueno", y nombres propios.
        4. ESTRUCTURA:
           - Inicia ÚNICAMENTE con: "Interviene el [CARGO] [NOMBRE], refiere que..."
           - El resto del texto debe ser casi idéntico al audio, solo ajustando la gramática a tercera persona.

        TEXTO A PROCESAR:
        "{texto_crudo}"
        """

        response = model.generate_content(prompt_usuario, safety_settings=safety_settings_nilo)
        return response.text
        
    except Exception as e:
        return f"Error con el modelo: {str(e)}"

def transcribir_audio(audio_file, api_key):
    aai.settings.api_key = api_key
    transcriber = aai.Transcriber()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
        tmp_file.write(audio_file.read())
        tmp_path = tmp_file.name
    config = aai.TranscriptionConfig(speaker_labels=True, language_code="es")
    try:
        transcript = transcriber.transcribe(tmp_path, config=config)
    except:
        return None
    os.remove(tmp_path)
    return transcript

def recuperar_por_id_debug(transcript_id, api_key):
    """
    Versión CORREGIDA para recuperar ID.
    Usa aai.Transcript.get_by_id() en lugar de transcriber.get_transcript()
    """
    aai.settings.api_key = api_key
    try:
        # --- CORRECCIÓN AQUÍ ---
        transcript = aai.Transcript.get_by_id(transcript_id)
        return transcript, None 
    except Exception as e:
        return None, str(e)

def dividir_texto_seguro(texto, tamano_max=12000):
    chunks = []
    while len(texto) > tamano_max:
        corte = texto.rfind('\n', 0, tamano_max)
        if corte == -1: corte = texto.rfind('. ', 0, tamano_max)
        if corte == -1: corte = tamano_max
        chunks.append(texto[:corte+1])
        texto = texto[corte+1:]
    chunks.append(texto)
    return chunks

# --- INTERFAZ ---
tab1, tab2 = st.tabs(["🎧 Audio / Recuperación", "📝 Texto Manual"])

# PESTAÑA 1: AUDIO
with tab1:
    # --- MÓDULO DE RECUPERACIÓN ---
    with st.container(border=True):
        st.subheader("📂 Recuperación de Sesión")
        st.info("Usa esto si ya subiste el audio pero se cerró la página. No gastará créditos.")
        
        col_rec1, col_rec2 = st.columns([3, 1])
        with col_rec1:
            id_recuperacion = st.text_input("Pega el ID de Transcripción aquí:", placeholder="Ej: 6m5p...")
        with col_rec2:
            st.write("") 
            st.write("") 
            boton_recuperar = st.button("♻️ Recuperar Ahora")
        
        if boton_recuperar:
            if not api_assembly:
                st.error("🚨 ERROR: El campo 'AssemblyAI API Key' en la barra lateral está vacío.")
            elif not id_recuperacion:
                st.error("⚠️ ERROR: No has escrito ningún ID.")
            else:
                id_limpio = id_recuperacion.strip()
                
                with st.spinner(f"Conectando con AssemblyAI para buscar ID: {id_limpio}..."):
                    recup, error_msg = recuperar_por_id_debug(id_limpio, api_assembly)
                    
                    if recup:
                        if recup.status == 'completed':
                            st.session_state.transcript_result = recup
                            st.success("✅ ¡Datos encontrados! Cargando...")
                            time.sleep(1)
                            st.rerun()
                        elif recup.status == 'processing':
                            st.warning("⏳ El audio aún se está procesando en la nube. Espera 1 minuto e intenta de nuevo.")
                        elif recup.status == 'error':
                            st.error("❌ El ID existe, pero la transcripción falló en el servidor de AssemblyAI.")
                    else:
                        st.error(f"❌ Error de conexión: {error_msg}")

    st.divider()

    uploaded_audio = st.file_uploader("O Sube un Audio Nuevo", type=["mp3", "wav", "m4a"])
    uploaded_template = st.file_uploader("Sube la Plantilla", type=["docx"], key="tpl1")
    
    if uploaded_audio and api_assembly and api_google:
        if "transcript_result" not in st.session_state:
            if st.button("🎙️ Iniciar Transcripción", key="btn_audio"):
                with st.spinner("Transcribiendo..."):
                    res = transcribir_audio(uploaded_audio, api_assembly)
                    if res:
                        st.session_state.transcript_result = res
                        st.rerun()

    if "transcript_result" in st.session_state:
        transcript = st.session_state.transcript_result
        if transcript:
            st.success(f"✅ Sesión Activa. ID: **{transcript.id}** (Cópialo para recuperar luego)")
            
            if transcript.utterances:
                unique_speakers = sorted(list(set([ut.speaker for ut in transcript.utterances])))
                st.subheader("Voces Detectadas")
                
                speaker_map = {}
                with st.form("mapping_form"):
                    for spk in unique_speakers:
                        sample = next((ut.text for ut in transcript.utterances if ut.speaker == spk), "")
                        st.markdown(f"**{spk}**: *'{sample[:100]}...'*")
                        c1, c2 = st.columns(2)
                        nombre = c1.text_input(f"Nombre ({spk})", placeholder="Ej: JUAN PEREZ")
                        cargo = c2.text_input(f"Cargo ({spk})", placeholder="Ej: Concejal", key=spk)
                        speaker_map[spk] = {"nombre": nombre, "cargo": cargo}
                        st.divider()
                    
                    if st.form_submit_button("🚀 Generar Acta"):
                        st.info("Procesando...")
                        
                        grouped_utterances = []
                        curr = {"speaker": transcript.utterances[0].speaker, "text": transcript.utterances[0].text}
                        for ut in transcript.utterances[1:]:
                            if ut.speaker == curr["speaker"]:
                                curr["text"] += " " + ut.text
                            else:
                                grouped_utterances.append(curr)
                                curr = {"speaker": ut.speaker, "text": ut.text}
                        grouped_utterances.append(curr)

                        doc_content = []
                        bar = st.progress(0)
                        for i, block in enumerate(grouped_utterances):
                            info = speaker_map.get(block["speaker"])
                            txt = redactar_bloque_gemini(block["text"], info["nombre"].upper(), info["cargo"], api_google, selected_model)
                            doc_content.append(txt)
                            time.sleep(2) # Pausa antibloqueo
                            bar.progress((i + 1) / len(grouped_utterances))
                        
                        st.session_state.final_text_audio = doc_content
                        st.success("¡Listo!")
            else:
                st.warning("⚠️ El audio fue procesado pero NO se detectaron diálogos. ¿El archivo tiene audio?")

            if "final_text_audio" in st.session_state:
                try:
                    doc = Document(uploaded_template) if uploaded_template else Document()
                except:
                    doc = Document()
                doc.add_paragraph("\n--- DESARROLLO ---\n")
                for p in st.session_state.final_text_audio:
                    doc.add_paragraph(p)
                    doc.add_paragraph("")
                
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                nombre_archivo = f"Acta_Final_{timestamp}.docx"
                doc.save(nombre_archivo)
                
                with open(nombre_archivo, "rb") as f:
                    st.download_button("📥 Descargar Word", f, nombre_archivo)

# PESTAÑA 2: TEXTO
with tab2:
    st.header("Texto Manual")
    if api_google:
        txt_input = st.text_area("Pega texto crudo aquí (Cualquier longitud):", height=400)
        c1, c2 = st.columns(2)
        n = c1.text_input("Nombre", "JUAN PEREZ", key="m_n")
        c = c2.text_input("Cargo", placeholder="Ej: Concejal", key="m_c")

        if st.button("Convertir"):
            if txt_input:
                with st.spinner("Procesando por bloques..."):
                    segmentos = dividir_texto_seguro(txt_input)
                    resultado_final = []
                    bar = st.progress(0)
                    
                    for i, seg in enumerate(segmentos):
                        res = redactar_bloque_gemini(seg, n.upper(), c, api_google, selected_model)
                        resultado_final.append(res)
                        bar.progress((i + 1) / len(segmentos))
                        time.sleep(2)
                    
                    st.session_state.res_manual = "\n\n".join(resultado_final)
                    st.success("¡Terminado!")
            else:
                st.error("Pega texto primero.")

        if "res_manual" in st.session_state:
            st.info("Resultado:")
            st.write(st.session_state.res_manual)
            
            doc_m = Document()
            for p in st.session_state.res_manual.split('\n'):
                if p.strip(): doc_m.add_paragraph(p)
            
            ts_m = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            nm = f"Acta_Manual_{ts_m}.docx"
            doc_m.save(nm)
            with open(nm, "rb") as f:
                st.download_button("📥 Descargar Word Manual", f, nm)